"""
Document Formatter Module
Handles the formatting of Word documents according to academic standards
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path
from datetime import datetime
import os


class DocumentFormatter:
    """Formats Word documents according to university academic standards"""
    
    def __init__(self):
        self.output_dir = Path("outputs")
        self.output_dir.mkdir(exist_ok=True)
        
        # University-specific formatting rules
        self.formatting_rules = {
            "National Standard": {
                "font_name": "Times New Roman",
                "font_size": 12,
                "line_spacing": 1.5,
                "margin_top": 1.18,  # inches
                "margin_bottom": 1.18,
                "margin_left": 1.57,
                "margin_right": 1.18,
                "paragraph_spacing_before": 0,
                "paragraph_spacing_after": 6,
            },
            "ITB": {
                "font_name": "Times New Roman",
                "font_size": 12,
                "line_spacing": 1.5,
                "margin_top": 1.18,
                "margin_bottom": 1.18,
                "margin_left": 1.57,
                "margin_right": 1.18,
                "paragraph_spacing_before": 0,
                "paragraph_spacing_after": 6,
            },
            "UI": {
                "font_name": "Times New Roman",
                "font_size": 12,
                "line_spacing": 2.0,
                "margin_top": 1.18,
                "margin_bottom": 1.18,
                "margin_left": 1.57,
                "margin_right": 1.18,
                "paragraph_spacing_before": 0,
                "paragraph_spacing_after": 0,
            },
            "UGM": {
                "font_name": "Times New Roman",
                "font_size": 12,
                "line_spacing": 1.5,
                "margin_top": 1.18,
                "margin_bottom": 1.18,
                "margin_left": 1.57,
                "margin_right": 1.18,
                "paragraph_spacing_before": 0,
                "paragraph_spacing_after": 6,
            }
        }
    
    def format_document(self, input_path: str, document_type: str, university: str) -> str:
        """
        Format a Word document according to specified university standards
        
        Args:
            input_path: Path to input document
            document_type: Type of document (thesis, paper, etc.)
            university: University template to use
            
        Returns:
            Path to formatted output document
        """
        # Load document
        doc = Document(input_path)
        
        # Get formatting rules
        rules = self.formatting_rules.get(university, self.formatting_rules["National Standard"])
        
        # Apply formatting
        self._apply_page_margins(doc, rules)
        self._apply_text_formatting(doc, rules)
        self._format_headings(doc, rules)
        self._fix_paragraph_spacing(doc, rules)
        
        # Generate output filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"formatted_{timestamp}.docx"
        output_path = self.output_dir / output_filename
        
        # Save formatted document
        doc.save(str(output_path))
        
        return str(output_path)
    
    def _apply_page_margins(self, doc: Document, rules: dict):
        """Apply page margins according to rules"""
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(rules["margin_top"])
            section.bottom_margin = Inches(rules["margin_bottom"])
            section.left_margin = Inches(rules["margin_left"])
            section.right_margin = Inches(rules["margin_right"])
    
    def _apply_text_formatting(self, doc: Document, rules: dict):
        """Apply text formatting to all paragraphs"""
        for paragraph in doc.paragraphs:
            # Skip empty paragraphs
            if not paragraph.text.strip():
                continue
            
            # Apply font and size
            for run in paragraph.runs:
                run.font.name = rules["font_name"]
                run.font.size = Pt(rules["font_size"])
            
            # Apply line spacing
            paragraph_format = paragraph.paragraph_format
            paragraph_format.line_spacing = rules["line_spacing"]
            
            # Apply alignment (justify for body text)
            if not paragraph.style.name.startswith('Heading'):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    def _format_headings(self, doc: Document, rules: dict):
        """Format heading styles"""
        # Get or create heading styles
        styles = doc.styles
        
        # Heading 1
        try:
            heading1 = styles['Heading 1']
            heading1.font.name = rules["font_name"]
            heading1.font.size = Pt(14)
            heading1.font.bold = True
            heading1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            heading1.paragraph_format.space_before = Pt(12)
            heading1.paragraph_format.space_after = Pt(6)
        except KeyError:
            pass
        
        # Heading 2
        try:
            heading2 = styles['Heading 2']
            heading2.font.name = rules["font_name"]
            heading2.font.size = Pt(13)
            heading2.font.bold = True
            heading2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            heading2.paragraph_format.space_before = Pt(12)
            heading2.paragraph_format.space_after = Pt(6)
        except KeyError:
            pass
        
        # Heading 3
        try:
            heading3 = styles['Heading 3']
            heading3.font.name = rules["font_name"]
            heading3.font.size = Pt(12)
            heading3.font.bold = True
            heading3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            heading3.paragraph_format.space_before = Pt(12)
            heading3.paragraph_format.space_after = Pt(6)
        except KeyError:
            pass
    
    def _fix_paragraph_spacing(self, doc: Document, rules: dict):
        """Fix paragraph spacing"""
        for paragraph in doc.paragraphs:
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(rules["paragraph_spacing_before"])
            paragraph_format.space_after = Pt(rules["paragraph_spacing_after"])
            
            # Set first line indent for body paragraphs
            if not paragraph.style.name.startswith('Heading'):
                paragraph_format.first_line_indent = Inches(0.5)
